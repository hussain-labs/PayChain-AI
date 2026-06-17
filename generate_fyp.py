import os
import time
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv('backend/.env')

API_KEY = os.getenv('GEMINI_API_KEY')
if not API_KEY:
    print("Error: GEMINI_API_KEY not found in backend/.env")
    exit(1)

genai.configure(api_key=API_KEY)

# Use standard gemini-pro or gemini-2.5-pro model
model = genai.GenerativeModel('gemini-2.5-pro')

def generate_content(prompt):
    print(f"Generating content for prompt: {prompt[:50]}...")
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error generating content: {e}")
        return "Content generation failed due to API error."

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    return p

doc = Document()

# Front Matter
title = doc.add_heading('Final Year Project Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\nPayChain-AI: A Web3 Cryptocurrency Dashboard with AI-driven Risk Analysis\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('BS(CS) (Session: 2021-25)\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

chapters = [
    {
        "title": "Chapter 1: Introduction",
        "subsections": [
            "1.1 Goals and objectives",
            "1.2 System statement of scope",
            "1.3 System context",
            "1.4 Theoretical Background (Blockchain, Web3, MERN Stack, AI Risk Analysis)",
            "1.5 Technology & Tools/hardware components"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive and lengthy Chapter 1 (Introduction) for a Final Year Project titled 'PayChain-AI', which is a Web3 crypto dashboard built with React, Node.js, Express, MongoDB, Wagmi/Viem, and Gemini AI. The chapter must cover: 1.1 Goals and objectives, 1.2 System statement of scope, 1.3 System context, 1.4 Theoretical Background (deep dive into Blockchain, Ethereum, Web3, MERN stack, and AI Risk Analysis), and 1.5 Technology & Tools. Write at least 2500 words. Do not use markdown headers, just output raw text with clear spacing."
    },
    {
        "title": "Chapter 2: Usage Scenario / User Interaction",
        "subsections": [
            "2.1 User profiles",
            "2.2 Use-cases",
            "2.3 Special usage considerations"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive and lengthy Chapter 2 (Usage Scenario) for the 'PayChain-AI' FYP report. It must cover: 2.1 User profiles (Admin, Regular User), 2.2 Use-cases (Wallet connection, Mock Transfers, AI Risk Analysis checking, Admin User Management, Support Ticketing), and 2.3 Special usage considerations. Write at least 2500 words."
    },
    {
        "title": "Chapter 3: Functional and Data Description",
        "subsections": [
            "3.1 System Architecture",
            "3.2 Data Description",
            "3.3 System Interface Description"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive and lengthy Chapter 3 (Functional and Data Description) for the 'PayChain-AI' FYP report. Detail the System Architecture (React frontend, Node/Express API, MongoDB database, Gemini AI integration, Web3 providers). Describe the Data Models (User Schema with isAdmin flag, SupportMessage Schema, Transaction histories). Describe System Interfaces (REST APIs, Web3 RPC calls, Gemini API). Write at least 2500 words."
    },
    {
        "title": "Chapter 4: Subsystem/module Description",
        "subsections": [
            "4.1 Web3 Integration Subsystem",
            "4.2 AI Risk Analysis Subsystem",
            "4.3 Admin Dashboard Subsystem",
            "4.4 Support Ticketing Subsystem"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive and lengthy Chapter 4 (Subsystem Description) for the 'PayChain-AI' FYP report. Detail every major module: Web3 Integration (Wagmi/Viem), AI Risk Analysis (Gemini prompt engineering and response parsing), Admin Dashboard (CRUD operations for users and tickets), and Support Ticketing. Provide flow descriptions, component details, restrictions, and performance issues. Write at least 2500 words."
    },
    {
        "title": "Chapter 5: Behavioral Model and Description",
        "subsections": [
            "5.1 Description for system behavior",
            "5.2 State Transition Diagrams Description",
            "5.3 Control specification"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive and lengthy Chapter 5 (Behavioral Model and Description) for the 'PayChain-AI' FYP report. Detail the events/interrupts (e.g., wallet disconnected, API failure, unauthorized admin access), the states (Idle, Authenticating, Fetching Data, Analyzing Risk), and the control specifications for JWT authentication and role-based access. Write at least 2000 words."
    },
    {
        "title": "Chapter 6: System Prototype Modeling and Simulation Results",
        "subsections": [
            "6.1 Description of system modeling approach",
            "6.2 Simulation results",
            "6.3 Special performance issues"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive Chapter 6 (System Prototype Modeling) for the 'PayChain-AI' FYP report. Discuss the agile prototyping approach used (MERN stack rapid development), simulation of crypto transfers without real gas fees (mock backend), latency of Gemini AI responses, and special performance issues with React state management and UI rendering. Write at least 2000 words."
    },
    {
        "title": "Chapter 7: System Estimates and Actual Outcome",
        "subsections": [
            "7.1 Historical data used for estimates",
            "7.2 Estimation techniques applied",
            "7.3 System Resources"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive Chapter 7 (System Estimates and Actual Outcome) for the 'PayChain-AI' FYP report. Detail software cost estimation techniques (like COCOMO II) applied to the MERN stack development, effort estimation in person-months, and the hardware/software resources required (Vite, MongoDB Atlas, Node.js environment, Gemini API tier). Write at least 1500 words."
    },
    {
        "title": "Chapter 8: Test Plan",
        "subsections": [
            "8.1 System Test and Procedure",
            "8.2 Testing strategy (Unit, Integration, Security)",
            "8.3 Testing tools and environment"
        ],
        "prompt": "You are an academic technical writer. Write an extremely comprehensive Chapter 8 (Test Plan) for the 'PayChain-AI' FYP report. Detail the unit testing of React components, integration testing of the Express API using Postman, Security testing (JWT validation, admin route protection, anti-XSS), and testing the Web3 connection states. Write at least 2000 words."
    },
    {
        "title": "Chapter 9 & 10: Future Enhancements and Conclusion",
        "subsections": [
            "Chapter 9: Future Enhancements and Recommendations",
            "Chapter 10: Conclusion / Summary"
        ],
        "prompt": "You are an academic technical writer. Write Chapter 9 and 10 for the 'PayChain-AI' FYP report. For Chapter 9, discuss future integration of actual Ethereum Smart Contracts for real token swaps, multi-chain support, and advanced AI agent execution. For Chapter 10, summarize the successful integration of Web3, AI, and traditional Web2 architecture in PayChain. Write at least 1500 words."
    }
]

print("Starting generation of 80-page FYP Document...")

for chapter in chapters:
    add_heading(doc, chapter['title'], level=1)
    for sub in chapter['subsections']:
        add_heading(doc, sub, level=2)
    
    # Generate content via Gemini
    content = generate_content(chapter['prompt'])
    add_paragraph(doc, content)
    doc.add_page_break()
    
    # Save incrementally
    doc.save('PayChain_FYP_Report.docx')
    print(f"Saved {chapter['title']}. Sleeping for 15 seconds to respect rate limits...")
    time.sleep(15)

print("Document generation complete! Saved as PayChain_FYP_Report.docx")
